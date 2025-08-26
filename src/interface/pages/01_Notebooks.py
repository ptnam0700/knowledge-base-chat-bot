from __future__ import annotations

import streamlit as st
from typing import List, Dict, Any
from datetime import datetime
import threading
import time
import re

from src.interface.notebooks import store
from src.interface.app_context import get_context
from src.utils.settings_manager import get_settings
from src.interface.notebooks.ingest import ingest_uploaded_files, ingest_url
from src.interface.utils.notebook_helper import NotebookHelper
import hashlib
import json
import os
from pathlib import Path
from src.utils.logger import logger
import streamlit.components.v1 as components
from PIL import Image
from src.interface.utils.prompt_text import (
    VI_CHAR_SET,
    MINDMAP_SUMMARY_VI,
    MINDMAP_JSON_INSTRUCTION_VI,
    NO_RESULTS_SYSTEM_VI,
    t,
    ts,
)

def _get_lang() -> str:
    try:
        persisted = get_settings()
        default_lang = persisted.get('language', 'vi')
    except Exception:
        default_lang = 'vi'
    # Prefer explicit language in session; fallback to app_settings, then default
    lang = st.session_state.get('language')
    if not lang and 'app_settings' in st.session_state:
        lang = (st.session_state.app_settings or {}).get('language')
    if not lang:
        lang = default_lang
    # Normalize and persist for consistency across pages
    try:
        st.session_state['language'] = lang
    except Exception:
        pass
    return lang

# Thread-safe in-memory task status (avoid using Streamlit APIs inside threads)
_TASK_STATUS: dict[str, dict] = {}

def _sanitize_for_filename(name: str) -> str:
    """Make a safe filename component from notebook name (ASCII-ish, underscores)."""
    try:
        import re
        safe = re.sub(r"\s+", "_", name.strip()) if name else "Notebook"
        safe = re.sub(r"[^\w\-]+", "", safe)
        return safe or "Notebook"
    except Exception:
        return "Notebook"

def _strip_markdown_for_docx(text: str) -> str:
    """Remove common Markdown decorations for DOCX plain text output."""
    try:
        import re
        if not text:
            return ""
        t = text
        # Inline code/backticks
        t = t.replace("```", "").replace("`", "")
        # Bold/italic markers
        t = t.replace("**", "").replace("__", "").replace("*", "")
        # Headings at line start like #, ##, ###
        t = re.sub(r"(?m)^\s*#{1,6}\s*", "", t)
        # Convert markdown links [text](url) -> text (url)
        t = re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", r"\1 (\2)", t)
        # Remove image syntax ![alt](url) -> alt (url)
        t = re.sub(r"!\[([^\]]*)\]\(([^\)]+)\)", r"\1 (\2)", t)
        # Normalize bullet prefixes to a simple dash
        t = re.sub(r"(?m)^\s*[-*â€¢]\s+", "- ", t)
        return t.strip()
    except Exception:
        return text or ""

def _get_notes_storage_path(notebook_id: str) -> Path:
    """Get the storage path for notes of a specific notebook."""
    data_dir = Path("data/notes")
    data_dir.mkdir(parents=True, exist_ok=True)
    return data_dir / f"notebook_{notebook_id}_notes.json"


def _save_notes_to_storage(notebook_id: str, notes: list):
    """Save notes to persistent storage."""
    try:
        storage_path = _get_notes_storage_path(notebook_id)
        with open(storage_path, 'w', encoding='utf-8') as f:
            json.dump(notes, f, ensure_ascii=False, indent=2, default=str)
    except Exception as e:
        st.error(f"Failed to save notes to storage: {e}")


def _load_notes_from_storage(notebook_id: str) -> list:
    """Load notes from persistent storage."""
    try:
        storage_path = _get_notes_storage_path(notebook_id)
        if storage_path.exists():
            with open(storage_path, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception as e:
        st.error(f"Failed to load notes from storage: {e}")
    return []


def _get_notebook_folder(notebook_id: str) -> Path:
    """Ensure and return the per-notebook working folder under data/notebooks/{id}."""
    root = Path("data/notebooks") / str(notebook_id)
    root.mkdir(parents=True, exist_ok=True)
    return root


def _get_latest_docx_path(notebook_id: str) -> Path:
    """Return latest DOCX report path if exists (Report_*.docx), fallback to legacy."""
    folder = _get_notebook_folder(notebook_id)
    latest = _find_latest_file(folder, ["Report_*.docx"]) or folder / "overview_latest.docx"
    return latest


def _get_latest_audio_path(notebook_id: str) -> Path:
    """Return latest MP3 report path if exists (Report_*.mp3), fallback to legacy."""
    folder = _get_notebook_folder(notebook_id)
    latest = _find_latest_file(folder, ["Report_*.mp3"]) or folder / "overview_latest.mp3"
    return latest


def _purge_old_overview_files(notebook_id: str, patterns: list[str], keep: list[Path]) -> None:
    """Delete old overview files matching patterns except those in keep list."""
    try:
        folder = _get_notebook_folder(notebook_id)
        keep_set = {p.resolve() for p in keep}
        for pat in patterns:
            for fp in folder.glob(pat):
                try:
                    if fp.resolve() not in keep_set:
                        fp.unlink()
                        logger.info(f"[StudioOverview] Purged old file: {fp.name}")
                except Exception:
                    continue
    except Exception:
        pass


def _get_latest_mindmap_path(notebook_id: str) -> Path:
    """Return stable path placeholder for future mindmap exports."""
    return _get_notebook_folder(notebook_id) / "mindmap_latest.png"

def _get_latest_mindmap_html_path(notebook_id: str) -> Path:
    """Return stable path for interactive mindmap HTML export."""
    return _get_notebook_folder(notebook_id) / "mindmap_latest.html"

def _is_valid_image(img_path: Path) -> bool:
    try:
        if not img_path.exists() or img_path.stat().st_size == 0:
            return False
        with Image.open(str(img_path)) as im:
            im.verify()
        return True
    except Exception:
        return False

def _extract_text_from_docx(docx_path: Path) -> str:
    """Extract plain text from a DOCX file. Fallback to empty string on error."""
    try:
        from docx import Document  # type: ignore
        doc = Document(str(docx_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception:
        return ""

def _generate_mindmap_outline(summary_text: str, *, prompt_manager=None, llm_client=None) -> dict:
    """Use LLM to build a hierarchical outline JSON for a mindmap.
    Schema: {"title": str, "nodes": [{"label": str, "children": [...] }]}
    Fallback to heuristic parsing if LLM unavailable or returns invalid JSON.
    """
    # Heuristic fallback
    def _heuristic_outline(text: str) -> dict:
        title = "Mindmap"
        try:
            first_line = next((ln.strip() for ln in (text or "").splitlines() if ln.strip()), "")
            if first_line:
                title = first_line[:80]
        except Exception:
            pass
        bullets = []
        for ln in (text or "").splitlines():
            s = ln.strip()
            if s.startswith(("- ", "* ", "â€¢ ")):
                bullets.append(s[2:].strip())
        children = [{"label": b, "children": []} for b in bullets[:20]]
        return {"title": title or "Mindmap", "nodes": children}

    if not summary_text or len(summary_text.strip()) < 10:
        return _heuristic_outline(summary_text)

    try:
        if prompt_manager and llm_client:
            instruction = MINDMAP_JSON_INSTRUCTION_VI
            messages = (
                prompt_manager.build_generic_prompt(
                    system_instruction=instruction,
                    user_content=summary_text[:12000],
                ) if hasattr(prompt_manager, "build_generic_prompt") else [
                    {"role": "system", "content": instruction},
                    {"role": "user", "content": summary_text[:12000]},
                ]
            )
            resp = llm_client.generate_response(messages, max_tokens=1200, temperature=0.1)
            content = resp.get("content", "") if isinstance(resp, dict) else getattr(resp, "content", "")
            import json as _json  # local import to avoid shadowing
            data = _json.loads(content)
            # basic validation
            if isinstance(data, dict) and "nodes" in data:
                return data
    except Exception:
        pass
    return _heuristic_outline(summary_text)

def _build_graphviz_from_outline(outline: dict):
    try:
        import graphviz  # type: ignore
        g = graphviz.Digraph(format="png")
        g.attr(rankdir="LR", nodesep="0.3", ranksep="0.5")
        g.node("root", outline.get("title", "Mindmap"), shape="box", style="rounded,filled", fillcolor="#EFF3FF")

        node_id_counter = {"v": 0}
        def add_nodes(parent_id: str, node: dict):
            node_id_counter["v"] += 1
            nid = f"n{node_id_counter['v']}"
            label = str(node.get("label", ""))[:80]
            g.node(nid, label, shape="box", style="rounded")
            g.edge(parent_id, nid)
            for ch in node.get("children", [])[:20]:
                add_nodes(nid, ch)

        for top in outline.get("nodes", [])[:30]:
            add_nodes("root", top)
        return g
    except Exception:
        return None

def _build_pyvis_from_outline(outline: dict, html_path: Path) -> bool:
    try:
        from pyvis.network import Network  # type: ignore
        net = Network(height="600px", width="100%", directed=False, bgcolor="#FFFFFF")
        net.barnes_hut()
        # root
        net.add_node("root", label=outline.get("title", "Mindmap"), shape="box", color="#AEC7E8")
        node_id_counter = {"v": 0}
        def add_nodes(parent_id: str, node: dict):
            node_id_counter["v"] += 1
            nid = f"n{node_id_counter['v']}"
            label = str(node.get("label", ""))[:80]
            net.add_node(nid, label=label, shape="box")
            net.add_edge(parent_id, nid)
            for ch in node.get("children", [])[:20]:
                add_nodes(nid, ch)
        for top in outline.get("nodes", [])[:60]:
            add_nodes("root", top)
        net.set_options('{"physics": {"stabilization": true}}')
        # Write HTML without trying to open a browser
        net.write_html(str(html_path), notebook=False)
        return html_path.exists()
    except Exception:
        return False

def _background_generate_mindmap(notebook_id: str, task_key: str) -> None:
    """Worker: build mindmap assets (interactive HTML + PNG fallback) and update status."""
    try:
        from src.interface.app_context import _build_context as _build_isolated_ctx  # type: ignore
        ctx = _build_isolated_ctx()
        nb = store.get_notebook(notebook_id)
        nb_name = getattr(nb, "name", "Notebook") if nb else "Notebook"

        # 1) Prefer existing DOCX for summary text
        docx_path = _get_latest_docx_path(notebook_id)
        summary_text = ""
        if docx_path.exists():
            summary_text = _extract_text_from_docx(docx_path)
        if not summary_text:
            # 2) Fallback: collect text and summarize via LLM
            full_text = _collect_notebook_texts(
                notebook_id,
                vector_db=ctx.get("vector_db"),
                search_engine=ctx.get("search_engine"),
            )
            summary_text = _run_langgraph_summary(
                full_text,
                prompt_manager=ctx.get("prompt_manager"),
                llm_client=ctx.get("llm_client"),
                additional_instructions=(
                    MINDMAP_SUMMARY_VI
                ),
                max_tokens=500,
            )

        # 3) Build outline
        outline = _generate_mindmap_outline(
            summary_text,
            prompt_manager=ctx.get("prompt_manager"),
            llm_client=ctx.get("llm_client"),
        )

        # 4) Persist outline
        nb_folder = _get_notebook_folder(notebook_id)
        outline_path = nb_folder / "mindmap_latest.json"
        try:
            with open(outline_path, "w", encoding="utf-8") as f:
                json.dump(outline, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

        # 5) Render interactive HTML first
        html_path = _get_latest_mindmap_html_path(notebook_id)
        html_ok = _build_pyvis_from_outline(outline, html_path)

        # 6) Fallback PNG via graphviz; also persist DOT for UI fallback
        png_path = _get_latest_mindmap_path(notebook_id)
        dot_path = _get_notebook_folder(notebook_id) / "mindmap_latest.dot"
        g = _build_graphviz_from_outline(outline)
        if g is not None:
            try:
                # Save DOT source for UI fallback
                with open(dot_path, "w", encoding="utf-8") as f:
                    f.write(g.source)
            except Exception:
                pass
        if not html_ok and g is not None:
            try:
                png_bytes = g.pipe(format="png")  # render to bytes
                with open(png_path, "wb") as f:
                    f.write(png_bytes)
            except Exception:
                # Leave only DOT if rendering fails (likely missing Graphviz `dot` binary)
                pass

        _TASK_STATUS[task_key] = {
            "running": False,
            "file_path": str(html_path if html_ok else png_path),
            "is_html": html_ok,
            "error": None,
        }
        logger.info(f"[Mindmap] Generated mindmap for notebook {notebook_id}")
    except Exception as e:
        _TASK_STATUS[task_key] = {"running": False, "file_path": None, "is_html": False, "error": str(e)}
        logger.error(f"[Mindmap] Failed: {e}")

def _collect_notebook_texts(notebook_id: str, *, vector_db=None, search_engine=None) -> str:
    """Collect and join all texts for a notebook from the metadata store."""
    try:
        if vector_db is None:
            ctx = get_context()
            vector_db = ctx.get("vector_db")
        if not vector_db:
            logger.warning(f"[StudioOverview] No vector_db available for notebook {notebook_id}")
            return ""
        metas = vector_db.metadata_manager.search_metadata({"notebook_id": notebook_id})
        logger.info(f"[StudioOverview] Collected {len(metas)} metadata entries for notebook {notebook_id}")
        texts = [m.get("text", "") for m in metas if m.get("text")]
        # Deduplicate and join with separators
        seen = set()
        unique_texts = []
        for t in texts:
            if t not in seen:
                unique_texts.append(t)
                seen.add(t)
        combined = "\n\n---\n\n".join(unique_texts)
        if combined.strip():
            logger.info(f"[StudioOverview] Combined text length: {len(combined)} chars (from metadata)")
            return combined

        # Fallback 1: load saved Studio notes from storage
        try:
            notes_path = _get_notes_storage_path(notebook_id)
            if notes_path.exists():
                notes = json.loads(notes_path.read_text(encoding="utf-8"))
                note_texts = [n.get("content", "") for n in notes if n.get("content")]
                if note_texts:
                    logger.info(f"[StudioOverview] Using {len(note_texts)} Studio notes as content source")
                    return "\n\n---\n\n".join(note_texts)
        except Exception:
            pass

        # Fallback 2: use cached overview from store
        try:
            ov = store.get_overview(notebook_id)
            if ov and ov.strip():
                logger.info("[StudioOverview] Using cached overview from store as content source")
                return ov
        except Exception:
            pass

        # Fallback: pull content via semantic search using broad queries
        if search_engine is None and vector_db is not None:
            try:
                from src.search.semantic_search import SemanticSearchEngine
                search_engine = SemanticSearchEngine(vector_db=vector_db)
            except Exception:
                search_engine = None
        if search_engine is None:
            logger.warning(f"[StudioOverview] No search engine available for notebook {notebook_id}")
            return ""
        queries = [
            "tÃ³m táº¯t ná»™i dung",
            "ná»™i dung chÃ­nh",
            "overview",
            "key points",
        ]
        collected = []
        for q in queries:
            try:
                results = search_engine.search(q, k=50, threshold=-1.0, filters={"notebook_id": notebook_id})
                logger.info(f"[StudioOverview] Search '{q}' returned {len(results)} results")
                for r in results:
                    txt = r.text or r.metadata.get("text", "")
                    if txt and txt not in collected:
                        collected.append(txt)
            except Exception:
                pass
        logger.info(f"[StudioOverview] Collected {len(collected)} snippets via search; combined length {sum(len(x) for x in collected)}")
        return "\n\n---\n\n".join(collected)
    except Exception:
        logger.error(f"[StudioOverview] Failed collecting texts for notebook {notebook_id}")
        return ""


def _find_latest_file(folder: Path, patterns: list[str]) -> Path | None:
    """Find the most recently modified file in folder matching any of patterns."""
    try:
        candidates: list[Path] = []
        for p in patterns:
            candidates.extend(sorted(folder.glob(p)))
        if not candidates:
            return None
        return max(candidates, key=lambda p: p.stat().st_mtime)
    except Exception:
        return None


def _run_langgraph_summary(
    full_text: str,
    *,
    prompt_manager=None,
    llm_client=None,
    additional_instructions: str | None = None,
    max_tokens: int | None = None,
) -> str:
    """Run LangGraph summarization workflow (with safe fallback) and return summary text."""
    if not full_text:
        return ""
    try:
        from src.ai.langgraph.workflows.summarization_workflow import create_summarization_workflow
    except Exception:
        create_summarization_workflow = None  # type: ignore
    try:
        if prompt_manager is None or llm_client is None:
            ctx = get_context()
            prompt_manager = ctx.get("prompt_manager")
            llm_client = ctx.get("llm_client")
        if not prompt_manager or not llm_client:
            return full_text[:50000]
        workflow = create_summarization_workflow(recursion_limit=8) if create_summarization_workflow else None
        # Determine instruction profile
        detailed_default = (
            "Viáº¿t báº£n tá»•ng quan CHI TIáº¾T, máº¡ch láº¡c báº±ng tiáº¿ng Viá»‡t (â‰ˆ800â€“1500 tá»«). "
            "Giá»¯ cáº¥u trÃºc má»™t bÃ¡o cÃ¡o hoÃ n chá»‰nh: (1) Má»Ÿ Ä‘áº§u ngáº¯n gá»n (1â€“3 cÃ¢u) nÃªu má»¥c tiÃªu/ná»™i dung chÃ­nh; "
            "(2) CÃ¡c má»¥c ná»™i dung theo chá»§ Ä‘á» (dÃ¹ng tiÃªu Ä‘á» ngáº¯n), dÆ°á»›i má»—i má»¥c liá»‡t kÃª bullets nÃªu Ã½ chÃ­nh, sá»‘ liá»‡u quan trá»ng, vÃ­ dá»¥ minh há»a; "
            "(3) Káº¿t luáº­n vÃ  khuyáº¿n nghá»‹/next steps. KhÃ´ng tÃ¡ch theo tá»«ng nguá»“n; há»£p nháº¥t thÃ´ng tin tá»« má»i nguá»“n thÃ nh má»™t bá»©c tranh thá»‘ng nháº¥t. "
            "Khi cÃ³ thá»ƒ, chÃ¨n trÃ­ch dáº«n ngáº¯n gá»n trong ngoáº·c vuÃ´ng [Nguá»“n: tiÃªu Ä‘á»/URL hoáº·c mÃ£ ghi chÃº] ngay sau dá»¯ kiá»‡n/sá»‘ liá»‡u. "
            "VÄƒn phong rÃµ rÃ ng, sÃºc tÃ­ch, chuyÃªn nghiá»‡p, Æ°u tiÃªn tÃ­nh chÃ­nh xÃ¡c."
        )
        instructions = additional_instructions or detailed_default
        is_concise = "150â€“250" in instructions or "150â€“250" in instructions
        gen_max_tokens = max_tokens if isinstance(max_tokens, int) and max_tokens > 0 else (600 if is_concise else 1600)
        if workflow is None:
            # Fallback: single-pass summary using prompt manager + llm_client
            messages = prompt_manager.build_summary_prompt(
                content=full_text[:50000],
                additional_instructions=instructions,
            )
            resp = llm_client.generate_response(messages, max_tokens=gen_max_tokens, temperature=0.2)
            content = resp.get("content", "") if isinstance(resp, dict) else getattr(resp, "content", "")
            return content
        state = workflow.invoke({
            "content": full_text,
            "llm_client": llm_client,
            "prompt_manager": prompt_manager,
            "remaining_loops": 2,
            "additional_instructions": instructions,
            "max_tokens": gen_max_tokens,
        })
        summary = state.get("summary", "")
        if not summary or not summary.strip():
            # Secondary fallback with stricter prompt
            messages = prompt_manager.build_summary_prompt(
                content=full_text[:50000],
                additional_instructions=instructions,
            )
            resp = llm_client.generate_response(messages, max_tokens=gen_max_tokens, temperature=0.2)
            return resp.get("content", "") if isinstance(resp, dict) else getattr(resp, "content", "")
        # Ensure final overview length aligns with instruction profile
        try:
            def _word_count(t: str) -> int:
                return len((t or "").split())
            if ((not is_concise and _word_count(summary) > 1100) or (is_concise and _word_count(summary) > 330)) and prompt_manager and llm_client:
                instruction = (
                    "RÃºt gá»n ná»™i dung sau thÃ nh "
                    + ("tá»•ng quan khoáº£ng 150â€“250 tá»«" if is_concise else "bÃ¡o cÃ¡o khoáº£ng 800â€“1500 tá»«")
                    + ", giá»¯ cáº¥u trÃºc máº¡ch láº¡c: má»Ÿ Ä‘áº§u ngáº¯n; cÃ¡c má»¥c theo chá»§ Ä‘á» vá»›i bullets nÃªu Ã½ chÃ­nh, sá»‘ liá»‡u, vÃ­ dá»¥; "
                    "káº¿t luáº­n/khuyáº¿n nghá»‹. KhÃ´ng tÃ¡ch theo nguá»“n; há»£p nháº¥t thÃ´ng tin; giá»¯ trÃ­ch dáº«n dáº¡ng [Nguá»“n: ...] náº¿u cÃ³."
                )
                messages = (
                    prompt_manager.build_generic_prompt(
                        system_instruction=instruction,
                        user_content=summary[:50000],
                    ) if hasattr(prompt_manager, "build_generic_prompt") else [
                        {"role": "system", "content": instruction},
                        {"role": "user", "content": summary[:50000]},
                    ]
                )
                resp = llm_client.generate_response(messages, max_tokens=gen_max_tokens, temperature=0.2)
                condensed = resp.get("content", "") if isinstance(resp, dict) else getattr(resp, "content", "")
                if condensed:
                    return condensed
        except Exception:
            pass
        return summary
    except Exception:
        return full_text[:50000]


def _background_generate_docx_overview(notebook_id: str, task_key: str) -> None:
    """Worker: generate DOCX overview and update session state when done."""
    try:
        # Build isolated context (avoid Streamlit session access in thread)
        from src.interface.app_context import _build_context as _build_isolated_ctx  # type: ignore
        ctx = _build_isolated_ctx()
        logger.info(f"[StudioOverview] (DOCX) Thread start for notebook {notebook_id}")
        # Gather content and summarize
        full_text = _collect_notebook_texts(
            notebook_id,
            vector_db=ctx.get("vector_db"),
            search_engine=ctx.get("search_engine"),
        )
        logger.info(f"[StudioOverview] (DOCX) Collected text length: {len(full_text)}")
        summary = _run_langgraph_summary(
            full_text,
            prompt_manager=ctx.get("prompt_manager"),
            llm_client=ctx.get("llm_client"),
            additional_instructions=(
                "Xuáº¥t báº£n bÃ¡o cÃ¡o DOCX: TrÃ¬nh bÃ y vÄƒn báº£n máº¡ch láº¡c, khÃ´ng dÃ¹ng Markdown (khÃ´ng **, __, #, ```). "
                "Sá»­ dá»¥ng Ä‘oáº¡n vÄƒn vÃ  dáº¥u gáº¡ch Ä‘áº§u dÃ²ng Ä‘Æ¡n giáº£n khi cáº§n; trÃ¡nh kÃ½ hiá»‡u Ä‘áº·c biá»‡t. "
                "Ná»™i dung chi tiáº¿t, rÃµ rÃ ng, chuyÃªn nghiá»‡p theo cáº¥u trÃºc bÃ¡o cÃ¡o."
            ),
        )
        if not summary or not summary.strip():
            logger.warning("[StudioOverview] (DOCX) Empty summary from LLM; falling back to raw text slice")
            summary = (full_text or "")

        # Create DOCX with timestamped Report_{name}_{datetime}.docx and update latest alias for UI
        nb_folder = _get_notebook_folder(notebook_id)
        try:
            nb = store.get_notebook(notebook_id)
            nb_name = getattr(nb, "name", "Notebook") if nb else "Notebook"
        except Exception:
            nb_name = "Notebook"
        safe_name = _sanitize_for_filename(nb_name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path: Path = nb_folder / f"Report_{safe_name}_{timestamp}.docx"
        try:
            from docx import Document  # type: ignore
            doc = Document()
            title = "Notebook Overview"
            try:
                nb = store.get_notebook(notebook_id)
                if nb and getattr(nb, "name", None):
                    title = f"Notebook Overview â€” {nb.name}"
            except Exception:
                pass
            doc.add_heading(title, level=1)
            if summary:
                doc.add_paragraph(_strip_markdown_for_docx(summary))
            else:
                doc.add_paragraph("No content available to summarize.")
            doc.save(str(file_path))
        except Exception:
            # Fallback to plain text file if python-docx is unavailable
            file_path = nb_folder / f"Report_{safe_name}_{timestamp}.txt"
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(_strip_markdown_for_docx(summary) if summary else "No content available to summarize.")

        # Purge legacy timestamped files by old pattern; keep new file
        _purge_old_overview_files(notebook_id, ["overview_*.docx", "overview_*.txt"], [file_path])

        # Update status (in-memory dict only)
        _TASK_STATUS[task_key] = {"running": False, "file_path": str(file_path), "error": None}
        logger.info(f"[StudioOverview] (DOCX) Written overview to {file_path}")
    except Exception as e:
        _TASK_STATUS[task_key] = {"running": False, "file_path": None, "error": str(e)}
        logger.error(f"[StudioOverview] (DOCX) Failed: {e}")


def _background_generate_audio_overview(notebook_id: str, task_key: str) -> None:
    """Worker: generate audio overview (MP3) and update session state when done."""
    try:
        # Build isolated context (avoid Streamlit session access in thread)
        from src.interface.app_context import _build_context as _build_isolated_ctx  # type: ignore
        ctx = _build_isolated_ctx()
        logger.info(f"[StudioOverview] (AUDIO) Thread start for notebook {notebook_id}")
        # Gather content and summarize
        full_text = _collect_notebook_texts(
            notebook_id,
            vector_db=ctx.get("vector_db"),
            search_engine=ctx.get("search_engine"),
        )
        logger.info(f"[StudioOverview] (AUDIO) Collected text length: {len(full_text)}")
        summary = _run_langgraph_summary(
            full_text,
            prompt_manager=ctx.get("prompt_manager"),
            llm_client=ctx.get("llm_client"),
        )
        if not summary or not summary.strip():
            logger.warning("[StudioOverview] (AUDIO) Empty summary from LLM; falling back to raw text slice")
            summary = (full_text or "")[:4000]

        tts_client = ctx.get("tts_client")
        voice = "alloy"
        model = "tts-1"
        if not tts_client:
            raise RuntimeError("TTS service not available")
        audio_bytes = tts_client.text_to_speech(
            text=summary[:4000] if summary else "No content available to summarize.",
            voice=voice,
            model=model,
            instructions="Speak clearly and at a moderate pace, suitable for an overview report."
        )
        if not audio_bytes:
            raise RuntimeError("Failed to synthesize audio")

        nb_folder = _get_notebook_folder(notebook_id)
        try:
            nb = store.get_notebook(notebook_id)
            nb_name = getattr(nb, "name", "Notebook") if nb else "Notebook"
        except Exception:
            nb_name = "Notebook"
        safe_name = _sanitize_for_filename(nb_name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = nb_folder / f"Report_{safe_name}_{timestamp}.mp3"
        with open(file_path, "wb") as f:
            f.write(audio_bytes)

        # Purge legacy timestamped files by old patterns; keep new file
        _purge_old_overview_files(notebook_id, ["overview_*.mp3", "overview_*.wav", "overview_*.m4a"], [file_path])

        _TASK_STATUS[task_key] = {"running": False, "file_path": str(file_path), "error": None}
        logger.info(f"[StudioOverview] (AUDIO) Written audio to {file_path}")
    except Exception as e:
        _TASK_STATUS[task_key] = {"running": False, "file_path": None, "error": str(e)}
        logger.error(f"[StudioOverview] (AUDIO) Failed: {e}")


def _generate_example_questions_from_summary(summary: str, *, prompt_manager=None, llm_client=None, target_lang: str = "vi") -> list[str]:
    """Generate 3 example questions using LLM based on the provided summary.

    Returns a list of up to 3 questions, short and diverse. Fallback to defaults if LLM fails.
    """
    try:
        if not prompt_manager or not llm_client or not summary:
            raise RuntimeError("Missing components or empty summary")
        if target_lang.startswith("vi"):
            instruction = (
                "Táº¡o Ä‘Ãºng 3 cÃ¢u há»i vÃ­ dá»¥ ngáº¯n gá»n, thÃ´ng minh, Ä‘a dáº¡ng kiá»ƒu (tÃ³m táº¯t, phÃ¢n tÃ­ch, so sÃ¡nh) "
                "dá»±a trÃªn pháº§n tÃ³m táº¯t dÆ°á»›i Ä‘Ã¢y. Tráº£ vá» má»—i cÃ¢u trÃªn má»™t dÃ²ng, khÃ´ng Ä‘Ã¡nh sá»‘, tiáº¿ng Viá»‡t."
            )
        else:
            instruction = (
                "Generate exactly 3 short, smart, diverse example questions (summary, analysis, comparison) "
                "based on the summary below. Return one question per line, no numbering, in English."
            )
        messages = prompt_manager.build_generic_prompt(
            system_instruction=instruction,
            user_content=summary[:6000],
        ) if hasattr(prompt_manager, "build_generic_prompt") else [
            {"role": "system", "content": instruction},
            {"role": "user", "content": summary[:6000]},
        ]
        resp = llm_client.generate_response(messages, max_tokens=256, temperature=0.3)
        content = resp.get("content", "") if isinstance(resp, dict) else getattr(resp, "content", "")
        lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
        # Filter out numbered bullets and take first 3
        cleaned = []
        for ln in lines:
            if ln[:2] in {"1.", "2.", "3."} or ln[:1] in {"-", "*", "â€¢"}:
                ln = ln[2:].strip() if ln[:2] in {"1.", "2.", "3."} else ln[1:].strip()
            if ln:
                cleaned.append(ln)
            if len(cleaned) >= 3:
                break
        if len(cleaned) >= 3:
            return cleaned[:3]
    except Exception:
        pass
    # Fallback
    return [
        "TÃ³m táº¯t 5 Ã½ chÃ­nh cá»§a tÃ i liá»‡u nÃ y, kÃ¨m trÃ­ch dáº«n nguá»“n.",
        "Liá»‡t kÃª má»‘c thá»i gian quan trá»ng vÃ  nguá»“n trÃ­ch dáº«n.",
        "So sÃ¡nh hai quan Ä‘iá»ƒm chÃ­nh trong cÃ¡c nguá»“n, kÃ¨m trÃ­ch dáº«n.",
    ]


def _background_generate_overview_and_examples(notebook_id: str, overview_task_key: str, examples_task_key: str) -> None:
    """Worker: generate text overview and example questions using LangChain without blocking UI."""
    try:
        from src.interface.app_context import _build_context as _build_isolated_ctx  # type: ignore
        ctx = _build_isolated_ctx()
        logger.info(f"[StudioOverview] (TEXT) Thread start for notebook {notebook_id}")
        full_text = _collect_notebook_texts(
            notebook_id,
            vector_db=ctx.get("vector_db"),
            search_engine=ctx.get("search_engine"),
        )
        logger.info(f"[StudioOverview] (TEXT) Collected text length: {len(full_text)}")

        # Generate overview via LangGraph/LLM
        summary = _run_langgraph_summary(
            full_text,
            prompt_manager=ctx.get("prompt_manager"),
            llm_client=ctx.get("llm_client"),
            additional_instructions=(
                "Viáº¿t má»™t báº£n tá»•ng quan ngáº¯n gá»n báº±ng tiáº¿ng Viá»‡t, Ä‘á»™ dÃ i khoáº£ng 150â€“250 tá»«. "
                "Cáº¥u trÃºc: 1â€“2 cÃ¢u má»Ÿ Ä‘áº§u ráº¥t ngáº¯n; sau Ä‘Ã³ cÃ¡c gáº¡ch Ä‘áº§u dÃ²ng nÃªu 4â€“6 Ã½ chÃ­nh; "
                "cÃ³ thá»ƒ thÃªm 1 Ä‘oáº¡n ngáº¯n káº¿t luáº­n náº¿u cáº§n. Ngáº¯n gá»n, rÃµ rÃ ng, trÃ¡nh chi tiáº¿t thá»«a."
            ),
            max_tokens=600,
        )
        if not summary or not summary.strip():
            nb = None
            try:
                nb = store.get_notebook(notebook_id)
            except Exception:
                pass
            summary = (getattr(nb, "description", None) or full_text or "").strip()[:4000]

        # Persist overview
        try:
            store.update_overview(notebook_id, summary)
        except Exception:
            pass
        _TASK_STATUS[overview_task_key] = {"running": False, "result": summary, "error": None}

        # Detect language from content to guide examples
        try:
            vi_chars = sum(ch in VI_CHAR_SET for ch in (full_text or "").lower())
            target_lang = "vi" if vi_chars > 0 else "en"
        except Exception:
            target_lang = "vi"

        # Generate examples via LLM
        questions = _generate_example_questions_from_summary(
            summary,
            prompt_manager=ctx.get("prompt_manager"),
            llm_client=ctx.get("llm_client"),
            target_lang=target_lang,
        )
        try:
            store.update_examples(notebook_id, questions)
        except Exception:
            pass
        _TASK_STATUS[examples_task_key] = {"running": False, "result": questions, "error": None}
        logger.info(f"[StudioOverview] (TEXT) Overview and examples generated for notebook {notebook_id}")
    except Exception as e:
        _TASK_STATUS[overview_task_key] = {"running": False, "result": None, "error": str(e)}
        _TASK_STATUS[examples_task_key] = {"running": False, "result": None, "error": str(e)}
        logger.error(f"[StudioOverview] (TEXT) Failed: {e}")

def _cleanup_old_overview_files(notebook_id: str, keep_latest: int = 1) -> dict[str, int]:
    """Clean up old overview files, keeping only the latest ones per type."""
    try:
        nb_folder = _get_notebook_folder(notebook_id)
        if not nb_folder.exists():
            return {"deleted": 0, "kept": 0}
        
        # Group files by type and find the latest ones
        file_groups = {
            "docx": [],
            "txt": [],
            "mp3": [],
            "wav": [],
            "m4a": []
        }
        
        for file_path in nb_folder.iterdir():
            if file_path.is_file():
                suffix = file_path.suffix.lower()
                if suffix in file_groups:
                    file_groups[suffix].append(file_path)
        
        deleted_count = 0
        kept_count = 0
        
        for file_type, files in file_groups.items():
            if len(files) <= keep_latest:
                kept_count += len(files)
                continue
                
            # Sort by modification time (newest first)
            sorted_files = sorted(files, key=lambda p: p.stat().st_mtime, reverse=True)
            
            # Keep the latest files
            files_to_keep = sorted_files[:keep_latest]
            files_to_delete = sorted_files[keep_latest:]
            
            # Delete old files
            for old_file in files_to_delete:
                try:
                    old_file.unlink()
                    deleted_count += 1
                    logger.info(f"[Cleanup] Deleted old {file_type} file: {old_file.name}")
                except Exception as e:
                    logger.warning(f"[Cleanup] Failed to delete {old_file.name}: {e}")
            
            kept_count += len(files_to_keep)
        
        logger.info(f"[Cleanup] Notebook {notebook_id}: deleted {deleted_count} old files, kept {kept_count} latest files")
        return {"deleted": deleted_count, "kept": kept_count}
        
    except Exception as e:
        logger.error(f"[Cleanup] Failed to cleanup notebook {notebook_id}: {e}")
        return {"deleted": 0, "kept": 0, "error": str(e)}


def _get_storage_info(notebook_id: str) -> dict[str, Any]:
    """Get storage information for a notebook's overview files."""
    try:
        nb_folder = _get_notebook_folder(notebook_id)
        if not nb_folder.exists():
            return {"total_size": 0, "file_count": 0, "file_types": {}}
        
        total_size = 0
        file_count = 0
        file_types = {}
        
        for file_path in nb_folder.iterdir():
            if file_path.is_file():
                try:
                    file_size = file_path.stat().st_size
                    total_size += file_size
                    file_count += 1
                    
                    suffix = file_path.suffix.lower()
                    if suffix not in file_types:
                        file_types[suffix] = {"count": 0, "size": 0}
                    file_types[suffix]["count"] += 1
                    file_types[suffix]["size"] += file_size
                except Exception:
                    pass
        
        return {
            "total_size": total_size,
            "file_count": file_count,
            "file_types": file_types,
            "folder_path": str(nb_folder)
        }
    except Exception as e:
        logger.error(f"[StorageInfo] Failed to get storage info for notebook {notebook_id}: {e}")
        return {"error": str(e)}


def _render_filters():
    return NotebookHelper.render_filters()

def _render_filters_and_sorting():
    """Render filters and sorting options in a single expandable section."""
    with st.expander(t("filter_sort_title", _get_lang()), expanded=False):
        # Initialize default sort option in session state
        if 'notebook_sort_option' not in st.session_state:
            st.session_state.notebook_sort_option = "date_new"
        
        # Check if sort option changed and clear cache if needed
        previous_sort = st.session_state.get('previous_sort_option', st.session_state.notebook_sort_option)
        if previous_sort != st.session_state.notebook_sort_option:
            # Clear all notebook caches when sort option changes
            for key in list(st.session_state.keys()):
                if key.startswith('notebooks_cache_'):
                    del st.session_state[key]
            st.session_state['previous_sort_option'] = st.session_state.notebook_sort_option
        
        # First row: Filters
        col1, col2, col3 = st.columns(3)
        
        with col1:
            q = st.text_input(t("search_by_name_desc_tag", _get_lang()), key="home_q", placeholder=t("search_by_name_desc_tag", _get_lang()))
            favorite_only = st.checkbox(t("favorites_only", _get_lang()), key="home_fav")
        
        with col2:
            date_from = st.date_input(t("date_from", _get_lang()), value=None, key="home_from")
            date_to = st.date_input(t("date_to", _get_lang()), value=None, key="home_to")
        
        with col3:
            st.write("")
            if st.button(t("clear_filters", _get_lang()), key="clear_filters"):
                st.session_state.pop("home_q", None)
                st.session_state.pop("home_fav", None)
                st.session_state.pop("home_from", None)
                st.session_state.pop("home_to", None)
                st.rerun()
        
        # Second row: Sorting options
        st.markdown("---")
        sort_col1, sort_col2 = st.columns([1, 1])
        
        with sort_col1:
            # Use stable internal keys; map to i18n labels with format_func
            sort_options = [
                "date_new",
                "date_old",
                "updated",
                "name_az",
                "name_za",
            ]
            sort_by = st.selectbox(
                t("sort_by", _get_lang()),
                options=sort_options,
                format_func=lambda k: {
                    "date_new": t("sort_option_date_new", _get_lang()),
                    "date_old": t("sort_option_date_old", _get_lang()),
                    "updated": t("sort_option_updated", _get_lang()),
                    "name_az": t("sort_option_name_az", _get_lang()),
                    "name_za": t("sort_option_name_za", _get_lang()),
                }.get(k, k),
                key="notebook_sort_option",
                help=t("sort_by", _get_lang())
            )
        
        with sort_col2:
            # Show current sorting method info
            if sort_by == "date_new":
                st.success(t("stable_sort_info_new", _get_lang()))
            elif sort_by == "date_old":
                st.success(t("stable_sort_info_old", _get_lang()))
            elif sort_by == "updated":
                st.warning(t("dynamic_sort_info", _get_lang()))
            else:
                st.info(t("alphabet_sort_info", _get_lang()))
        
        # Add tip about sorting
        st.caption(t("alphabet_sort_info", _get_lang()))
    
    return q, favorite_only, date_from.isoformat() if date_from else None, date_to.isoformat() if date_to else None, sort_by


def _get_notebooks_info_for_llm():
    """Get information about all notebooks for LLM context."""
    try:
        notebooks = store.list_notebooks()
        notebooks_info = []
        
        for nb in notebooks[:10]:  # Limit to 10 notebooks to avoid context overflow
            nb_info = {
                'id': nb.id,
                'name': nb.name,
                'description': nb.description or "No description",
                'tags': nb.tags or [],
                'sources_count': len(nb.sources) if nb.sources else 0,
                'created_at': nb.created_at.split('T')[0] if nb.created_at else 'Unknown',
                'is_favorite': nb.is_favorite
            }
            notebooks_info.append(nb_info)
        
        return notebooks_info
    except Exception as e:
        logger.error(f"Failed to get notebooks info: {e}")
        return []


def _generate_no_results_response(query: str, notebook_name: str, notebooks_info: list) -> str:
    """Generate a helpful response when no relevant content is found."""
    try:
        # Detect language
        is_vietnamese = any(ch in VI_CHAR_SET for ch in query.lower())
        
        # Try to use LLM for intelligent response generation
        try:
            ctx = get_context()
            if ctx.get("prompt_manager") and ctx.get("llm_client") and notebooks_info:
                # Create context about available notebooks
                notebooks_context = []
                for nb in notebooks_info[:5]:  # Limit to 5 for context
                    nb_desc = f"- {nb['name']}: {nb['description']} (Tags: {', '.join(nb['tags'][:3]) if nb['tags'] else 'None'}, Sources: {nb['sources_count']})"
                    notebooks_context.append(nb_desc)
                
                context_text = "\n".join(notebooks_context)
                
                if is_vietnamese:
                    system_prompt = f"""Báº¡n lÃ  trá»£ lÃ½ AI há»¯u Ã­ch. NgÆ°á»i dÃ¹ng vá»«a tÃ¬m kiáº¿m cÃ¢u há»i '{query}' trong notebook '{notebook_name}' nhÆ°ng khÃ´ng tÃ¬m tháº¥y káº¿t quáº£.

HÃ£y táº¡o pháº£n há»“i há»¯u Ã­ch báº±ng tiáº¿ng Viá»‡t bao gá»“m:
1. ThÃ´ng bÃ¡o khÃ´ng tÃ¬m tháº¥y ná»™i dung liÃªn quan
2. Gá»£i Ã½ cÃ¡c notebook phÃ¹ há»£p tá»« danh sÃ¡ch dÆ°á»›i Ä‘Ã¢y (náº¿u cÃ³)
3. HÆ°á»›ng dáº«n ngÆ°á»i dÃ¹ng cÃ³ thá»ƒ tÃ¬m kiáº¿m trong notebook khÃ¡c hoáº·c táº¡o notebook má»›i thá»§ cÃ´ng

LÆ¯U Ã QUAN TRá»ŒNG: KHÃ”NG Ä‘Æ°á»£c há»©a háº¹n hoáº·c gá»£i Ã½ ráº±ng báº¡n cÃ³ thá»ƒ tá»± Ä‘á»™ng táº¡o notebook. Chá»‰ hÆ°á»›ng dáº«n há» lÃ m thá»§ cÃ´ng.

Danh sÃ¡ch notebooks hiá»‡n cÃ³:
{context_text}

HÃ£y tráº£ lá»i ngáº¯n gá»n, há»¯u Ã­ch vÃ  thÃ¢n thiá»‡n."""
                else:
                    system_prompt = f"""You are a helpful AI assistant. The user just searched for '{query}' in notebook '{notebook_name}' but no relevant content was found.

Please generate a helpful response in English that includes:
1. Acknowledgment that no relevant content was found
2. Suggest relevant notebooks from the list below (if any)
3. Guide the user to search in other notebooks or manually create a new notebook

IMPORTANT NOTE: DO NOT promise or suggest that you can automatically create notebooks. Only guide them to do it manually.

Available notebooks:
{context_text}

Keep the response concise, helpful, and friendly."""
                
                messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Query: '{query}' in notebook '{notebook_name}'"}
                ]
                
                response = ctx["llm_client"].generate_response(
                    messages,
                    max_tokens=300,
                    temperature=0.3
                )
                
                llm_response = response.get("content") if isinstance(response, dict) else getattr(response, "content", "")
                if llm_response and len(llm_response.strip()) > 50:
                    return llm_response
        except Exception as e:
            logger.warning(f"Failed to generate LLM response for no results: {e}")
        
        # Fallback to static response
        if is_vietnamese:
            response = f"âŒ **KhÃ´ng tÃ¬m tháº¥y ná»™i dung liÃªn quan**\n\n"
            response += f"TÃ´i khÃ´ng thá»ƒ tÃ¬m tháº¥y thÃ´ng tin vá» **'{query}'** trong notebook **'{notebook_name}'**.\n\n"
            
            if notebooks_info:
                response += "**ðŸ’¡ Gá»£i Ã½ tÃ¬m kiáº¿m:**\n"
                response += "Báº¡n cÃ³ thá»ƒ thá»­ tÃ¬m kiáº¿m trong cÃ¡c notebook khÃ¡c:\n\n"
                
                # Show relevant notebooks based on query
                relevant_notebooks = []
                for nb in notebooks_info:
                    relevance_score = 0
                    if query.lower() in nb['name'].lower():
                        relevance_score += 3
                    if query.lower() in nb['description'].lower():
                        relevance_score += 2
                    if any(query.lower() in tag.lower() for tag in nb['tags']):
                        relevance_score += 1
                    
                    if relevance_score > 0:
                        relevant_notebooks.append((nb, relevance_score))
                
                # Sort by relevance and show top 3
                relevant_notebooks.sort(key=lambda x: x[1], reverse=True)
                for nb, score in relevant_notebooks[:3]:
                    response += f"â€¢ **{nb['name']}** ({nb['sources_count']} sources, {nb['created_at']})\n"
                    if nb['description'] != "No description":
                        response += f"  _{nb['description'][:100]}{'...' if len(nb['description']) > 100 else ''}_\n"
                
                response += "\n**ðŸ†• Hoáº·c táº¡o notebook má»›i:**\n"
                response += "Náº¿u khÃ´ng cÃ³ notebook nÃ o phÃ¹ há»£p, báº¡n cÃ³ thá»ƒ táº¡o notebook má»›i thá»§ cÃ´ng tá»« trang chÃ­nh vá»›i ná»™i dung liÃªn quan Ä‘áº¿n cÃ¢u há»i cá»§a mÃ¬nh."
            else:
                response += "**ðŸ’¡ Gá»£i Ã½:** Báº¡n cÃ³ thá»ƒ táº¡o notebook má»›i thá»§ cÃ´ng tá»« trang chÃ­nh rá»“i thá»­ láº¡i"
        else:
            response = f"âŒ **No relevant content found**\n\n"
            response += f"I couldn't find information about **'{query}'** in notebook **'{notebook_name}'**.\n\n"
            
            if notebooks_info:
                response += "**ðŸ’¡ Search suggestions:**\n"
                response += "You can try searching in other notebooks:\n\n"
                
                # Show relevant notebooks based on query
                relevant_notebooks = []
                for nb in notebooks_info:
                    relevance_score = 0
                    if query.lower() in nb['name'].lower():
                        relevance_score += 3
                    if query.lower() in nb['description'].lower():
                        relevance_score += 2
                    if any(query.lower() in tag.lower() for tag in nb['tags']):
                        relevance_score += 1
                    
                    if relevance_score > 0:
                        relevant_notebooks.append((nb, relevance_score))
                
                # Sort by relevance and show top 3
                relevant_notebooks.sort(key=lambda x: x[1], reverse=True)
                for nb, score in relevant_notebooks[:3]:
                    response += f"â€¢ **{nb['name']}** ({nb['sources_count']} sources, {nb['created_at']})\n"
                    if nb['description'] != "No description":
                        response += f"  _{nb['description'][:100]}{'...' if len(nb['description']) > 100 else ''}_\n"
                
                response += "\n**ðŸ†• Or create a new notebook:**\n"
                response += "If no notebooks are suitable, you can manually create a new notebook from the main page with content related to your question."
            else:
                response += "**ðŸ’¡ Suggestion:** You can manually create a new notebook and try again"
        
        return response
        
    except Exception as e:
        logger.error(f"Failed to generate no results response: {e}")
        # Fallback response
        return f"âŒ KhÃ´ng tÃ¬m tháº¥y ná»™i dung liÃªn quan Ä‘áº¿n '{query}' trong notebook '{notebook_name}'. Vui lÃ²ng thá»­ tÃ¬m kiáº¿m trong notebook khÃ¡c hoáº·c táº¡o notebook má»›i thá»§ cÃ´ng tá»« trang chÃ­nh."


def _render_notebook_card(nb: store.Notebook, col):
    return NotebookHelper.render_notebook_card(nb, col)


def _note_title_from_content(content: str, max_len: int = 60) -> str:
    """Generate a concise human-friendly title from note content."""
    if not content:
        return "Untitled"
    try:
        text = content.strip()
        # Remove fenced code blocks
        text = re.sub(r"```[\s\S]*?```", " ", text)
        # Remove inline code backticks
        text = re.sub(r"`([^`]*)`", r"\1", text)
        # Replace links/images with their visible text
        text = re.sub(r"!\[([^\]]*)\]\([^\)]*\)", r"\1", text)  # images -> alt text
        text = re.sub(r"\[([^\]]+)\]\([^\)]*\)", r"\1", text)    # links -> label
        # Strip markdown headings prefixes
        text = re.sub(r"^\s*#+\s*", "", text, flags=re.MULTILINE)
        # First non-empty line
        first_line = next((ln.strip() for ln in text.splitlines() if ln.strip()), "")
        candidate = re.sub(r"\s+", " ", first_line) if first_line else re.sub(r"\s+", " ", text)
        candidate = candidate.strip()
        # Remove emphasis markers like **bold**, *italic*, __bold__, _italic_
        candidate = re.sub(r"\*\*([^*]+)\*\*", r"\1", candidate)
        candidate = re.sub(r"\*([^*]+)\*", r"\1", candidate)
        candidate = re.sub(r"__([^_]+)__", r"\1", candidate)
        candidate = re.sub(r"_([^_]+)_", r"\1", candidate)
        # Clean leftover multiple asterisks
        candidate = re.sub(r"\*{2,}", "", candidate)
        if len(candidate) > max_len:
            cutoff = candidate.rfind(" ", 0, max_len)
            if cutoff == -1:
                cutoff = max_len
            candidate = candidate[:cutoff].rstrip() + "..."
        return candidate or "Untitled"
    except Exception:
        return "Untitled"


def _render_sources_panel(nb: store.Notebook):
    """Render sources panel."""
    st.subheader(t("sources", _get_lang()))

    # Ensure Studio notes are available for better note titles
    if 'studio_notes' not in st.session_state:
        st.session_state.studio_notes = {}
    if nb.id not in st.session_state.studio_notes:
        st.session_state.studio_notes[nb.id] = []
    if not st.session_state.studio_notes[nb.id]:
        try:
            loaded_notes = _load_notes_from_storage(nb.id)
            if loaded_notes:
                st.session_state.studio_notes[nb.id] = loaded_notes
        except Exception:
            pass
    
    # Lightweight cached helpers to avoid repeated IO on reruns
    @st.cache_data(show_spinner=False, ttl=3600)
    def _cached_page_title(url: str) -> str | None:
        try:
            return get_context()["document_processor"].get_page_title(url)
        except Exception:
            return None

    @st.cache_data(show_spinner=False, ttl=3600)
    def _cached_youtube_title(url: str) -> str | None:
        try:
            yinfo = get_context()["youtube_processor"].get_video_info(url)
            return yinfo.get("title")
        except Exception:
            return None

    @st.cache_data(show_spinner=False, ttl=300)
    def _cached_meta_search(source: str, notebook_id: str) -> list[dict]:
        try:
            ctx = get_context()
            return ctx['vector_db'].metadata_manager.search_metadata({
                'source': source,
                'notebook_id': notebook_id,
            }) if ctx.get('vector_db') else []
        except Exception:
            return []

    # Sources list
    if not nb.sources:
        st.info(t("no_sources", _get_lang()))
    else:
        # Scrollable container for long source lists
        st.markdown('<div style="max-height:200px; overflow-y:auto; padding-right:8px;">', unsafe_allow_html=True)
        # Prepare metadata manager for chunk counts
        try:
            ctx = get_context()
            metadata_manager = ctx.get("vector_db").metadata_manager if ctx.get("vector_db") else None
        except Exception:
            metadata_manager = None

        for i, s in enumerate(nb.sources):
            col1, col2 = st.columns([4, 1])
            with col1:
                # Prefer Studio-like generated titles for note sources
                display_title = s.title
                display_href = None
                try:
                    if getattr(s, 'type', '') == 'note':
                        notes_list = st.session_state.studio_notes.get(nb.id, [])
                        matched = None
                        # Try by meta.note_id
                        meta = getattr(s, 'meta', None) or {}
                        note_id = meta.get('note_id') if isinstance(meta, dict) else None
                        if note_id:
                            for n in notes_list:
                                if n.get('id') == note_id or n.get('original_chat_id') == note_id:
                                    matched = n
                                    break
                        # Fallback by content prefix that was stored as source_path_or_url
                        if matched is None:
                            for n in notes_list:
                                prefix = (n.get('content') or "")[:100] + "..."
                                if getattr(s, 'source_path_or_url', None) == prefix:
                                    matched = n
                                    break
                        if matched is not None:
                            display_title = _note_title_from_content(matched.get('content', ''))
                        else:
                            # Fallback: try derive from stored snippet
                            display_title = _note_title_from_content(getattr(s, 'source_path_or_url', '') or s.title)
                    elif getattr(s, 'type', '') in ['url', 'youtube']:
                        # Attempt to fetch a human-friendly page title
                        url_value = getattr(s, 'source_path_or_url', s.title)
                        display_href = url_value
                        page_title = None
                        # Try using metadata stored in vector DB for this source (cached)
                        matches = _cached_meta_search(url_value, nb.id)
                        if matches:
                            meta0 = matches[0]
                            page_title = meta0.get('youtube_title') or meta0.get('page_title')
                        # Fallback: for youtube use processor to get title on-demand
                        if page_title is None and s.type == 'youtube':
                            page_title = _cached_youtube_title(url_value)
                        # Fallback: try fetch <title> from URL for generic web pages
                        if page_title is None and s.type == 'url':
                            page_title = _cached_page_title(url_value)
                        if page_title:
                            display_title = page_title
                        else:
                            display_title = url_value
                except Exception:
                    pass
                # Render as hyperlink if we have an href
                header_text = f"[{display_title}]({display_href})" if display_href else display_title
                with st.expander(f"{header_text}", expanded=False):
                    # Build rich details JSON for the source
                    chunk_count = 0
                    content_type = s.type
                    try:
                        if metadata_manager is not None:
                            src_key = getattr(s, 'source_path_or_url', s.title)
                            matches = _cached_meta_search(src_key, nb.id)
                            chunk_count = len(matches)
                            if matches:
                                content_type = matches[0].get('content_type', content_type)
                    except Exception:
                        pass
                    st.json({
                        'source': getattr(s, 'source_path_or_url', s.title),
                        'content_type': content_type,
                        'chunk_count': chunk_count,
                        'added_at': s.added_at,
                    })
                    if s.meta:
                        st.json(s.meta)
            with col2:
                if st.button("âŒ", key=f"del_source_{s.id}", help=t("delete_source", _get_lang())):
                    store.remove_source(nb.id, s.id)
                    st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("---")
    lang = _get_lang()
    st.subheader(t("add_sources", lang))
    uploaded = st.file_uploader(
        t("upload_files", lang),
        type=["mp4","avi","mov","mkv","mp3","wav","pdf","docx","txt","xlsx"],
        accept_multiple_files=True,
        key="nb_upload",
    )
    url = st.text_input(
        t("or_add_link", lang),
        key="nb_url",
        placeholder=t("enter_urls_placeholder", lang),
    )
    # Internet search UI (keys are namespaced per-notebook to avoid cross-notebook leakage)
    st.markdown(t("search_internet", lang))
    with st.container():
        search_col, btn_col = st.columns([5,1])
        with search_col:
            search_query = st.text_input(
                t("search_internet", lang),
                key=f"nb_web_search_{nb.id}",
                max_chars=100,
                placeholder=t("enter_keywords_placeholder", lang),
                label_visibility="collapsed",
            )
        with btn_col:
            do_search = st.button(t("search", lang), key=f"nb_web_search_btn_{nb.id}", use_container_width=True)
    # Persist search results in session
    results_key = f"nb_web_results_{nb.id}"
    if results_key not in st.session_state:
        st.session_state[results_key] = []
    if do_search and search_query and search_query.strip():
        try:
            ctx = get_context()
            urls = ctx["web_search"].web_search(search_query.strip())
            st.session_state[results_key] = urls
            if not urls:
                st.info(t("no_results_for_query", lang))
        except Exception as e:
            st.error(f"{t('web_search_failed', lang)}: {e}")
            st.session_state[results_key] = []
    # Render selectable results
    if st.session_state.get(results_key):
        st.markdown("#### " + t("search_internet", lang))
        select_all = st.checkbox(t("select_all", lang), key=f"nb_web_select_all_{nb.id}", value=False)
        selected_urls = []
        for idx, link in enumerate(st.session_state[results_key][:20]):
            checked = st.checkbox(link, key=f"nb_web_result_{nb.id}_{idx}", value=select_all)
            if checked:
                selected_urls.append(link)
    else:
        selected_urls = []
    if st.button(t("add_to_notebook", lang), key="nb_add_btn"):
        added = 0
        duplicates: list[str] = []
        # Build existing set of identifiers to avoid duplicates
        existing_keys = set()
        for s in nb.sources or []:
            key = (s.type, getattr(s, 'source_path_or_url', s.title))
            existing_keys.add(key)
        if uploaded:
            added += ingest_uploaded_files(nb.id, uploaded)
            for f in uploaded:
                key = ("file", f.name)
                if key in existing_keys:
                    duplicates.append(f.name)
                    continue
                store.add_source(nb.id, type="file", title=f.name, source_path_or_url=f.name)
                existing_keys.add(key)
        if url:
            src_type = ("youtube" if ("youtube.com" in url or "youtu.be" in url) else "url")
            key = (src_type, url)
            if key in existing_keys:
                duplicates.append(url)
            else:
                added += ingest_url(nb.id, url)
                # Try to resolve a better title for url/youtube
                resolved_title = url
                try:
                    if src_type == 'youtube':
                        yinfo = get_context()["youtube_processor"].get_video_info(url)
                        resolved_title = yinfo.get("title") or url
                    else:
                        page_title = get_context()["document_processor"].get_page_title(url)
                        resolved_title = page_title or url
                except Exception:
                    pass
                store.add_source(nb.id, type=src_type, title=resolved_title, source_path_or_url=url)
                existing_keys.add(key)
        # Add selected web search links
        for weblink in selected_urls:
            src_type = ("youtube" if ("youtube.com" in weblink or "youtu.be" in weblink) else "url")
            key = (src_type, weblink)
            if key in existing_keys:
                duplicates.append(weblink)
                continue
            try:
                resolved_title = weblink
                if src_type == 'youtube':
                    yinfo = get_context()["youtube_processor"].get_video_info(weblink)
                    resolved_title = yinfo.get("title") or weblink
                else:
                    page_title = get_context()["document_processor"].get_page_title(weblink)
                    resolved_title = page_title or weblink
                added += ingest_url(nb.id, weblink)
                store.add_source(nb.id, type=src_type, title=resolved_title, source_path_or_url=weblink)
                existing_keys.add(key)
            except Exception:
                continue
        if duplicates and added == 0:
            if len(duplicates) == 1:
                st.warning(t("source_exists", lang))
            else:
                st.warning(t("some_sources_exist", lang))
        elif duplicates:
            st.warning(t("some_sources_skipped", lang) + ", ".join(duplicates[:3]) + ("..." if len(duplicates) > 3 else ""))
        # Clear web search results and query after any Add action to avoid repetition
        try:
            st.session_state[results_key] = []
            st.session_state[f"nb_web_search_{nb.id}"] = ""
            st.session_state[f"nb_web_select_all_{nb.id}"] = False
        except Exception:
            pass
        if added > 0:
            st.success(t("added_chunks", lang).format(n=added))
            st.rerun()


def _notebook_overview(nb: store.Notebook):
    with st.expander(t("overview_examples", _get_lang()), expanded=False):
        st.markdown("### " + t("overview", _get_lang()))

        # Check if we need to regenerate overview and examples
        current_sources_count = len(nb.sources)
        cached_sources_count = st.session_state.get(f"sources_count_{nb.id}", 0)
        stored_overview = store.get_overview(nb.id)
        stored_examples = nb.examples
        
        # Only regenerate if sources count changed or no cached data exists
        needs_regeneration = (
            current_sources_count != cached_sources_count or 
            not stored_overview or 
            not stored_examples
        )
        
        overview_task_key = f"studio_text_overview_task_{nb.id}"
        examples_task_key = f"studio_text_examples_task_{nb.id}"
        if overview_task_key not in _TASK_STATUS:
            _TASK_STATUS[overview_task_key] = {"running": False, "result": None, "error": None}
        if examples_task_key not in _TASK_STATUS:
            _TASK_STATUS[examples_task_key] = {"running": False, "result": None, "error": None}

        if needs_regeneration:
            # Trigger background generation once; UI remains responsive
            if not _TASK_STATUS[overview_task_key]["running"] and not _TASK_STATUS[examples_task_key]["running"]:
                _TASK_STATUS[overview_task_key] = {"running": True, "result": None, "error": None}
                _TASK_STATUS[examples_task_key] = {"running": True, "result": None, "error": None}
                threading.Thread(
                    target=_background_generate_overview_and_examples,
                    args=(nb.id, overview_task_key, examples_task_key),
                    daemon=True,
                ).start()
            # Show placeholders while background tasks run
            overview = stored_overview or t("creating_overview", _get_lang())
            examples = stored_examples or [t("creating_examples", _get_lang())]
            # Update cached sources count so we don't retrigger until change
            st.session_state[f"sources_count_{nb.id}"] = current_sources_count
        else:
            # Use cached data
            overview = stored_overview
            examples = stored_examples
        
        st.write(overview)
        
        # Show cache status
        if not needs_regeneration:
            st.caption(t("cached_overview_examples", _get_lang()))

        st.markdown("#### " + t("examples", _get_lang()))
        cols = st.columns(len(examples))
        for i, ex in enumerate(examples[:3]):
            with cols[i]:
                if st.button(ex, key=f"ex_{i}"):
                    st.session_state["nb_chat_input"] = ex


def _generate_example_questions(nb: store.Notebook) -> List[str]:
    return NotebookHelper.generate_example_questions(nb)


def _render_chat(nb: store.Notebook):
    """Render chat interface with conversation view and save note per answer."""
    ctx = get_context()

    # Initialize chat histories per-notebook in session to avoid cross-notebook leakage
    if 'chat_histories' not in st.session_state:
        st.session_state.chat_histories = {}
    if nb.id not in st.session_state.chat_histories:
        st.session_state.chat_histories[nb.id] = []  # list of {id, question, answer, timestamp, sources}
    chat_history = st.session_state.chat_histories[nb.id]
    
    # Initialize button counter to ensure unique keys
    if 'chat_button_counter' not in st.session_state:
        st.session_state.chat_button_counter = 0

    # Preference for including sources (used when rendering history)
    include_sources_pref = st.session_state.get('nb_include_sources', True)

    # Ensure Studio notes store exists for this notebook
    if 'studio_notes' not in st.session_state:
        st.session_state.studio_notes = {}
    if nb.id not in st.session_state.studio_notes:
        st.session_state.studio_notes[nb.id] = []

    # Status placeholder
    status_ph = st.empty()

    # Show searching status when searching (avoid heavy re-render while toggling UI)
    if st.session_state.get('is_searching', False):
        status_ph.info(t("searching_answering", _get_lang()))
    else:
        # Chat history display area - only show when not searching
        st.markdown("### " + t("chat_history", _get_lang()))
        if chat_history:
            st.info(f"ðŸ“ {t('chat_history', _get_lang())}: {len(chat_history)} messages")
            
            # Chat history styling
            from src.interface.utils.notebook_ui import NotebookUI
            chat_style = NotebookUI.chat_style_css(max_height=500)
            st.markdown(chat_style, unsafe_allow_html=True)
            st.markdown('<div class="nb-chat-wrapper">', unsafe_allow_html=True)
            
            for item in chat_history[-50:]:
                st.markdown(f'<div class="nb-chat-item nb-chat-q"><div class="bubble q-bubble">{item["question"]}</div></div>', unsafe_allow_html=True)
                st.markdown(f'<div class="nb-chat-item nb-chat-a"><div class="bubble a-bubble">{item["answer"]}</div></div>', unsafe_allow_html=True)
                
                # Action buttons for each message - now with 3 columns to accommodate Speak button
                bcols = st.columns([1,1,1,3])
                with bcols[0]:
                    # Stable key per chat message to avoid key drift across reruns
                    button_key = f"save_note_{item['id']}"
                    
                    # Check if this note was already saved
                    # Use per-notebook storage for studio notes
                    notes_for_this_nb = st.session_state.studio_notes.get(nb.id, [])
                    note_already_saved = any(n.get('original_chat_id') == item['id'] for n in notes_for_this_nb)
                    
                    if st.button(t("save_note", _get_lang()) if not note_already_saved else t("saved", _get_lang()), 
                               key=button_key, 
                               disabled=note_already_saved):
                        # Ensure per-notebook notes store exists
                        if 'studio_notes' not in st.session_state:
                            st.session_state.studio_notes = {}
                        if nb.id not in st.session_state.studio_notes:
                            st.session_state.studio_notes[nb.id] = []
                        note = {
                            'id': f"note_{item['id']}",
                            'original_chat_id': item['id'],  # Track original chat message
                            'content': item['answer'],
                            'timestamp': datetime.now().isoformat(),
                            'sources': item.get('sources', []),
                            'question': item.get('question', ''),
                            'added_to_source': False
                        }
                        st.session_state.studio_notes[nb.id].append(note)
                        
                        # Save to persistent storage
                        _save_notes_to_storage(nb.id, st.session_state.studio_notes[nb.id])
                        
                        st.success(t("saved", _get_lang()))
                        st.rerun()  # Rerun to update button state
                
                with bcols[1]:
                    # Speak button for text-to-speech
                    if st.button(t("speak", _get_lang()), key=f"speak_chat_{item['id']}", 
                               help=t("listen_answer", _get_lang()), use_container_width=True):
                        try:
                            ctx = get_context()
                            if ctx.get("tts_client"):
                                with st.spinner(t("audio_generating", _get_lang())):
                                    # Get answer content for TTS
                                    answer_text = item['answer']
                                    
                                    # Truncate text if too long (TTS has limits)
                                    max_length = 4000
                                    if len(answer_text) > max_length:
                                        answer_text = answer_text[:max_length] + "..."
                                        st.warning(t("answer_truncated", _get_lang()).format(n=max_length))
                                    
                                    # Generate audio using TTS client
                                    audio_data = ctx["tts_client"].text_to_speech(
                                        text=answer_text,
                                        voice="alloy",  # Fixed voice
                                        model="tts-1",  # Fixed model - you can change this
                                        instructions="Speak clearly and at a moderate pace, suitable for answer reading."
                                    )
                                    
                                    if audio_data:
                                        # Create audio player with autoplay
                                        st.audio(audio_data, format="audio/mp3", start_time=0)
                                        st.success(t("audio_generated", _get_lang()))
                                    else:
                                        st.error(t("audio_failed", _get_lang()))
                            else:
                                st.error(t("tts_not_available", _get_lang()))
                        except Exception as e:
                            st.error(f"{t('error_generating_speech', _get_lang())}: {str(e)}")
                
                with bcols[2]:
                    if include_sources_pref and item.get('sources'):
                        with st.popover("Sources"):
                            for s in item['sources'][:5]:
                                st.markdown(f"- {s}")
            
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.info("ðŸ’¬ " + t("no_chat_history", _get_lang()))

    st.markdown("---")
    
    # Auto scroll to this section if flag is set
    if st.session_state.get('auto_scroll_to_question', False):
        st.session_state['auto_scroll_to_question'] = False
        # Add JavaScript for smooth scrolling
        from src.interface.utils.notebook_ui import NotebookUI
        st.markdown(NotebookUI.smooth_scroll_to_question_js(), unsafe_allow_html=True)
    
    st.markdown("### " + t("ask_question", _get_lang()))

    # Group small controls together; put fast_mode first to minimize layout shifts
    col1, col2, col3 = st.columns([1,1,1])
    with col1:
        fast_mode = st.checkbox(t("fast_mode", _get_lang()), value=False, help="Faster answers with smaller k and shorter generations")
    with col2:
        use_cot = st.checkbox(t("cot", _get_lang()), value=False)
    with col3:
        include_src = st.checkbox(t("include_sources", _get_lang()), value=include_sources_pref)
        st.session_state['nb_include_sources'] = include_src
        
    # Input + Ask in the same container and row
    with st.container(key="ask_question"):
        if st.session_state.get('clear_nb_input', False):
            st.session_state['nb_chat_input'] = ""
            st.session_state['clear_nb_input'] = False
        # Put a single label across both columns for better alignment
        input_col, ask_col = st.columns([9,1])
        with input_col:
            query = st.text_area(t("ask_question", _get_lang()), key="nb_chat_input", placeholder=t("your_question_placeholder", _get_lang()), label_visibility="collapsed")
        with ask_col:
            ask_clicked = st.button(t("ask", _get_lang()), type="primary", use_container_width=True)

    # Ask button handler
    if ask_clicked:
        if not query.strip():
            st.warning(t("please_enter_question", _get_lang()))
            return
        
        # Set searching flag to hide chat history and show status
        st.session_state['is_searching'] = True
        
        try:
            # Smaller top-k in fast mode for speed
            k_value = 4 if fast_mode else 8
            # Use a stricter threshold to avoid off-topic retrieval
            search_threshold = 0.35
            results = ctx["search_engine"].search(query, k=k_value, threshold=search_threshold, filters={"notebook_id": nb.id})
            # Extra guard: filter any items below threshold just in case
            filtered_results = [r for r in (results or []) if (getattr(r, "score", 0) or 0) >= search_threshold]
            if not filtered_results:
                # Generate helpful response when no results found
                notebooks_info = _get_notebooks_info_for_llm()
                no_results_answer = _generate_no_results_response(query, nb.name, notebooks_info)
                
                # Create a chat message for no results
                new_message = {
                    'id': f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}",
                    'question': query,
                    'answer': no_results_answer,
                    'timestamp': datetime.now().isoformat(),
                    'sources': [],
                    'structured_output': None
                }
                
                # Append to this notebook's chat history
                st.session_state.chat_histories[nb.id].append(new_message)
                
                # Show success message
                st.success(t("helpful_response_generated", _get_lang()) + " Scroll up to see the chat history.")
                
                # Reset searching flag to show chat history
                st.session_state['is_searching'] = False
                
                # Set flags for UI behavior on rerun
                st.session_state['auto_scroll_to_question'] = True
                st.session_state['clear_nb_input'] = True
                
                # Force rerun to show updated chat history
                st.rerun()
            
            rag_results = [
                {"text": r.text or r.metadata.get("text", ""), "score": r.score, "metadata": r.metadata}
                for r in filtered_results
            ]
            # Detect question language and enforce consistent answer language
            is_vietnamese = any(ch in VI_CHAR_SET for ch in query.lower())
            lang_instruction = (
                "HÃ£y tráº£ lá»i báº±ng tiáº¿ng Viá»‡t, vÄƒn phong rÃµ rÃ ng, máº¡ch láº¡c."
                if is_vietnamese else
                "Answer in English with clear and concise language."
            )
            role_instruction = (
                "Báº¡n pháº£i luÃ´n tráº£ lá»i báº±ng tiáº¿ng Viá»‡t."
                if is_vietnamese else
                "You must always respond in English."
            )
            # Prepare context: skip explicit summarization in fast mode to save time
            rag_for_sum = rag_results[:5] if fast_mode else rag_results
            if fast_mode:
                summary_text = "\n\n".join([r.get("text", "") for r in rag_for_sum])[:4000]
            else:
                # Use LangChain prompt + LLM to summarize when not in fast mode
                if ctx.get("prompt_manager") and ctx.get("llm_client"):
                    messages = ctx["prompt_manager"].build_summary_prompt(
                        content="\n\n".join([r.get("text", "") for r in rag_for_sum]),
                        additional_instructions=lang_instruction,
                    )
                    summary_resp = ctx["llm_client"].generate_response(
                        messages,
                        max_tokens=600,
                    )
                    summary_text = summary_resp.get("content", "")
                elif ctx.get("legacy_summarizer"):
                    summary_result = ctx["legacy_summarizer"].summarize_chunks(
                        rag_for_sum,
                        query=query,
                        use_chain_of_thought=use_cot,
                        extra_instructions=lang_instruction,
                        role=role_instruction,
                    )
                    summary_text = summary_result.summary

            # Use LangChain as primary AI system
            from src.utils.feature_flags import feature_flags

            answer = summary_text
            structured_output = None
            try:
                # Build prompt using LangChain prompt manager
                messages = ctx["prompt_manager"].build_qa_prompt(
                    context=summary_text,
                    question=query,
                    use_chain_of_thought=use_cot and not fast_mode,
                )
                messages.insert(0, {"role": "system", "content": role_instruction})

                # Use LangChain LLM client
                response = ctx["llm_client"].generate_response(
                    messages,
                    use_memory=not fast_mode,
                    store_in_memory=not fast_mode,
                    max_memory_context=3,
                    context_sources=[r.metadata.get("source","unknown") for r in filtered_results],
                    # Per-call overrides for latency in fast mode
                    max_tokens=600 if fast_mode else None,
                )
                # Normalize response format
                answer = response.get("content") if isinstance(response, dict) else getattr(response, "content", answer)
                
                # Parse structured output if available
                if feature_flags.is_enabled("use_structured_output") and ctx.get("output_parser"):
                    try:
                        structured_output = ctx["output_parser"].parse_qa(answer)
                        sources = [r.metadata.get("source", "unknown") for r in filtered_results[:5]]
                        structured_output.citations = ctx["output_parser"].extract_citations(answer, sources)
                    except Exception:
                        pass
            except Exception as _e:
                # Fallback to legacy system if LangChain fails
                try:
                    if ctx.get("legacy_prompt_engineer") and ctx.get("legacy_llm_client"):
                        pe = ctx["legacy_prompt_engineer"]
                        messages = pe.build_qa_prompt(
                            context=summary_text,
                            question=query,
                            use_chain_of_thought=use_cot and not fast_mode,
                        )
                        messages.insert(0, {"role": "system", "content": role_instruction})
                        
                        response = ctx["legacy_llm_client"].generate_response(
                            messages,
                            use_memory=not fast_mode,
                            store_in_memory=not fast_mode,
                            max_memory_context=3,
                            context_sources=[r.metadata.get("source","unknown") for r in filtered_results],
                        )
                        answer = response.content
                except Exception:
                    # Keep summary_text as final fallback
                    pass

            # Append new message to chat history with structured output
            new_message = {
                'id': f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}",
                'question': query,
                'answer': answer,
                'timestamp': datetime.now().isoformat(),
                'sources': [r.metadata.get('source','unknown') for r in filtered_results[:5]],
                'structured_output': structured_output
            }
            # Append to this notebook's chat history only
            st.session_state.chat_histories[nb.id].append(new_message)
            
            # Show success message
            st.success(t("answer_generated", _get_lang()) + " Scroll up to see the chat history.")
            
            # Reset searching flag to show chat history
            st.session_state['is_searching'] = False
            
            # Set flags for UI behavior on rerun
            st.session_state['auto_scroll_to_question'] = True
            st.session_state['clear_nb_input'] = True
            
            # Force rerun to show updated chat history
            st.rerun()
            
        except Exception as e:
            st.error(f"Error generating answer: {str(e)}")
            # Reset searching flag to show chat history after error
            st.session_state['is_searching'] = False


def _render_studio_actions(nb: store.Notebook, layout: str | None = None):
    """Render action buttons.
    layout: 'docx_only' | 'audio_only' | None (both rows previously).
    """
    task_docx_key = f"studio_docx_task_{nb.id}"
    task_audio_key = f"studio_audio_task_{nb.id}"

    # Initialize task states
    if task_docx_key not in _TASK_STATUS:
        _TASK_STATUS[task_docx_key] = {"running": False, "file_path": None, "error": None}
    if task_audio_key not in _TASK_STATUS:
        _TASK_STATUS[task_audio_key] = {"running": False, "file_path": None, "error": None}

    if layout in (None, "docx_only"):
        # DOCX button (not full width)
        docx_state = _TASK_STATUS.get(task_docx_key, {"running": False})
        docx_label = t("generating_docx", _get_lang()) if docx_state.get("running") else t("btn_docx", _get_lang())
        if st.button(docx_label, key=f"btn_docx_{nb.id}", disabled=docx_state.get("running", False)):
            _TASK_STATUS[task_docx_key] = {"running": True, "file_path": None, "error": None}
            threading.Thread(target=_background_generate_docx_overview, args=(nb.id, task_docx_key), daemon=True).start()
        if docx_state.get("error"):
            st.error(f"{t('error', _get_lang())}: {docx_state['error']}")

    if layout in (None, "audio_only"):
        # AUDIO button (not full width)
        audio_state = _TASK_STATUS.get(task_audio_key, {"running": False})
        audio_label = t("generating_audio", _get_lang()) if audio_state.get("running") else t("btn_audio", _get_lang())
        if st.button(audio_label, key=f"btn_audio_{nb.id}", disabled=audio_state.get("running", False)):
            _TASK_STATUS[task_audio_key] = {"running": True, "file_path": None, "error": None}
            threading.Thread(target=_background_generate_audio_overview, args=(nb.id, task_audio_key), daemon=True).start()
        if audio_state.get("error"):
            st.error(f"{t('error', _get_lang())}: {audio_state['error']}")


def _render_studio_panel(nb: store.Notebook):
    """Render studio panel with saved notes (no file manager UI)."""

    # Initialize per-notebook Studio notes store
    if 'studio_notes' not in st.session_state:
        st.session_state.studio_notes = {}
    if nb.id not in st.session_state.studio_notes:
        st.session_state.studio_notes[nb.id] = []
    
    # Load notes from persistent storage if session state is empty
    if not st.session_state.studio_notes[nb.id]:
        stored_notes = _load_notes_from_storage(nb.id)
        if stored_notes:
            st.session_state.studio_notes[nb.id] = stored_notes
            st.success(f"ðŸ“š Loaded {len(stored_notes)} notes from storage")

    # One-time migration from legacy saved_notes list (if present)
    if 'saved_notes' in st.session_state and st.session_state.saved_notes:
        st.session_state.studio_notes[nb.id].extend(st.session_state.saved_notes)
        st.session_state.saved_notes = []
        # Save migrated notes to storage
        _save_notes_to_storage(nb.id, st.session_state.studio_notes[nb.id])

    notes = st.session_state.studio_notes[nb.id]
    # Deduplicate notes by original_chat_id (fallback id)
    try:
        unique_notes_map = {}
        for n in notes:
            dedup_key = n.get('original_chat_id') or n.get('id')
            if dedup_key not in unique_notes_map:
                unique_notes_map[dedup_key] = n
        if len(unique_notes_map) != len(notes):
            st.session_state.studio_notes[nb.id] = list(unique_notes_map.values())
            notes = st.session_state.studio_notes[nb.id]
            # Save deduplicated notes to storage
            _save_notes_to_storage(nb.id, notes)
    except Exception:
        pass

    # Sync Studio notes with current sources: if a previously added note's source was deleted,
    # re-enable the Add to Source button by resetting flags.
    try:
        existing_source_ids = {s.id for s in (nb.sources or [])}
        changed = False
        for n in notes:
            if not n.get('added_to_source'):
                continue
            src_id = n.get('source_id')
            # If we have a stored source_id but it no longer exists, try to relink or reset
            if src_id and src_id not in existing_source_ids:
                matched = None
                # Try to find by meta.note_id if available
                note_identifier = n.get('id') or n.get('original_chat_id')
                for s in (nb.sources or []):
                    meta = getattr(s, 'meta', None) or {}
                    if isinstance(meta, dict) and meta.get('note_id') == note_identifier:
                        matched = s
                        break
                # Fallback: match by content prefix used when creating source
                if matched is None:
                    prefix = (n.get('content') or "")[:100] + "..."
                    for s in (nb.sources or []):
                        try:
                            if s.type == 'note' and s.source_path_or_url == prefix:
                                matched = s
                                break
                        except Exception:
                            continue
                if matched is not None:
                    n['source_id'] = matched.id
                    n['added_to_source'] = True
                    changed = True
                else:
                    n['added_to_source'] = False
                    n.pop('source_id', None)
                    changed = True
            # If we do not have a stored source_id but note is marked added, try to verify existence
            elif not src_id:
                matched = None
                note_identifier = n.get('id') or n.get('original_chat_id')
                for s in (nb.sources or []):
                    meta = getattr(s, 'meta', None) or {}
                    if isinstance(meta, dict) and meta.get('note_id') == note_identifier:
                        matched = s
                        break
                if matched is None:
                    prefix = (n.get('content') or "")[:100] + "..."
                    for s in (nb.sources or []):
                        try:
                            if s.type == 'note' and s.source_path_or_url == prefix:
                                matched = s
                                break
                        except Exception:
                            continue
                if matched is not None:
                    n['source_id'] = matched.id
                    n['added_to_source'] = True
                    changed = True
                else:
                    n['added_to_source'] = False
                    changed = True
        if changed:
            _save_notes_to_storage(nb.id, notes)
    except Exception:
        pass

    if not notes:
        st.info(t("no_saved_notes", _get_lang()))
        return
    
    # Display saved notes
    for i, note in enumerate(notes):
        generated_name = _note_title_from_content(note.get('content', ''))
        with st.expander(f"ðŸ“ {generated_name} ({t('note_word', _get_lang())} {i+1} - {note['timestamp'][:16]})", expanded=False):
            st.markdown(note['content'])
            st.caption(f"{t('sources_label', _get_lang())}: {', '.join(note['sources'][:2])}")
            
            # Action buttons - now with 3 columns to accommodate Speak button
            col1, col2, col3 = st.columns([2, 1, 1])
            
            with col1:
                add_label = t("add_to_source", _get_lang()) if not note.get('added_to_source') else t("added", _get_lang())
                if st.button(add_label, key=f"add_to_source_{note['id']}_{i}", 
                           help="Convert note to source", disabled=note.get('added_to_source', False)):
                    # Add note as a new source
                    # Compose a better source title from content, fallback to date
                    generated_name = _note_title_from_content(note.get('content', ''))
                    title_text = generated_name or f"Note from {note['timestamp'][:10]}"
                    sref = store.add_source(
                        nb.id, 
                        type="note", 
                        title=title_text, 
                        source_path_or_url=note['content'][:100] + "...",
                        meta={
                            "origin": "studio_note",
                            "note_id": note.get('id') or note.get('original_chat_id')
                        }
                    )
                    # Mark note as added to sources but keep it in Studio
                    note['added_to_source'] = True
                    if sref is not None:
                        note['source_id'] = sref.id
                    # Save updated notes to storage
                    _save_notes_to_storage(nb.id, st.session_state.studio_notes[nb.id])
                    st.success(t("note_added_to_sources", _get_lang()) + " (kept in Studio)")
                    st.rerun()
            
            with col2:
                # Speak button for text-to-speech
                if st.button(t("speak", _get_lang()), key=f"speak_note_{note['id']}_{i}", 
                           help=t("listen_answer", _get_lang()), use_container_width=True):
                    try:
                        # Get context and TTS client with fallback
                        ctx = get_context()
                        tts_client = ctx.get("tts_client")
                        
                        # If TTS client is None, try to re-initialize it
                        if tts_client is None:
                            st.info("ðŸ”„ Re-initializing TTS client...")
                            try:
                                from src.ai.tts_client import TTSClient
                                tts_client = TTSClient()
                                # Update context
                                ctx["tts_client"] = tts_client
                                st.success("âœ… TTS client re-initialized!")
                            except Exception as init_error:
                                st.error(f"âŒ Failed to re-initialize TTS client: {init_error}")
                                return
                        
                        if tts_client:
                            with st.spinner("ðŸŽµ Generating audio..."):
                                # Get note content for TTS
                                note_text = note['content']
                                    
                                # Truncate text if too long (TTS has limits)
                                max_length = 4000
                                if len(note_text) > max_length:
                                    note_text = note_text[:max_length] + "..."
                                    st.warning(f"âš ï¸ Note was truncated for TTS (max {max_length} characters)")
                                    
                                # Generate audio using TTS client with fixed model
                                audio_data = tts_client.text_to_speech(
                                    text=note_text,
                                    voice="alloy",  # Fixed voice
                                    model="tts-1",  # Fixed model - you can change this
                                    instructions="Speak clearly and at a moderate pace, suitable for note reading."
                                )
                                
                                if audio_data:
                                    # Create audio player with autoplay
                                    st.audio(audio_data, format="audio/mp3", start_time=0)
                                    st.success(t("audio_generated", _get_lang()))
                                else:
                                    st.error(t("audio_failed", _get_lang()))
                        else:
                            st.error(t("tts_not_available", _get_lang()))
                    except Exception as e:
                        st.error(f"{t('error_generating_speech', _get_lang())}: {str(e)}")
                        st.info("ðŸ’¡ This might be due to API rate limits or network issues")
            
            with col3:
                if st.button(t("delete", _get_lang()), key=f"delete_note_{note['id']}_{i}", 
                           help=t("delete", _get_lang())):
                    st.session_state.studio_notes[nb.id].pop(i)
                    # Save updated notes to storage
                    _save_notes_to_storage(nb.id, st.session_state.studio_notes[nb.id])
                    st.success(t("note_deleted", _get_lang()))
                    st.rerun()


def _render_create_view():
    edit_notebook_id = st.session_state.get("edit_notebook_id")
    is_edit_mode = edit_notebook_id is not None

    if is_edit_mode:
        st.title(t("edit_notebook", _get_lang()))
        notebook = store.get_notebook(edit_notebook_id)
        if not notebook:
            st.error("Notebook not found")
            st.query_params["view"] = "list"
            st.rerun()
            return
    else:
        st.title(t("create_notebook", _get_lang()))

    with st.form("create_notebook_form"):
        name = st.text_input(
            t("field_notebook_name", _get_lang()),
            placeholder="e.g., Marketing Q4 Reports",
            value=notebook.name if is_edit_mode else ""
        )
        desc = st.text_area(
            t("field_description_optional", _get_lang()),
            value=notebook.description if is_edit_mode else ""
        )
        tags = st.text_input(
            t("field_tags", _get_lang()),
            value=", ".join(notebook.tags) if is_edit_mode and notebook.tags else ""
        )
        uploaded = st.file_uploader(
            t("field_upload_files", _get_lang()),
            type=["mp4","avi","mov","mkv","mp3","wav","pdf","docx","txt","xlsx"],
            accept_multiple_files=True,
        )
        url = st.text_input(t("field_add_link", _get_lang()))
        submitted = st.form_submit_button(t("btn_save", _get_lang()) if is_edit_mode else t("btn_create", _get_lang()), type="primary")

    if submitted:
        if not name.strip():
            st.error(t("error_name_required", _get_lang()))
            return

        if is_edit_mode:
            store.update_notebook(
                edit_notebook_id,
                name=name.strip(),
                description=desc.strip(),
                tags=[t.strip() for t in tags.split(',') if t.strip()]
            )
            st.success(t("msg_notebook_updated", _get_lang()))
            st.session_state.pop("edit_notebook_id", None)
            st.session_state["current_notebook_id"] = edit_notebook_id
            st.query_params["view"] = "notebook"
            st.rerun()
        else:
            nb = store.create_notebook(name=name.strip(), description=desc.strip(), tags=[t.strip() for t in tags.split(',') if t.strip()])
            added = 0
            if uploaded:
                added += ingest_uploaded_files(nb.id, uploaded)
                for f in uploaded:
                    store.add_source(nb.id, type="file", title=f.name, source_path_or_url=f.name)
            if url:
                added += ingest_url(nb.id, url)
                store.add_source(nb.id, type=("youtube" if ("youtube.com" in url or "youtu.be" in url) else "url"), title=url, source_path_or_url=url)

            st.success(t("msg_notebook_created", _get_lang()).format(n=added))
            st.session_state["current_notebook_id"] = nb.id
            st.query_params["view"] = "notebook"
            st.rerun()

    if st.button(t("back_to_notebooks", _get_lang())):
        st.session_state.pop("edit_notebook_id", None)
        st.query_params["view"] = "list"
        st.rerun()


def main():
    st.set_page_config(page_title="Notebooks", page_icon="ðŸ““", layout="wide")
    
    # Load custom CSS
    with open("src/interface/styles/notebook_layout.css") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

    action = st.query_params.get("action", None)
    notebook_id = st.query_params.get("id", None)
    view = st.query_params.get("view", "list")

    if action == "open" and notebook_id:
        st.session_state["current_notebook_id"] = notebook_id
        st.query_params["view"] = "notebook"
        st.rerun()
    elif action == "favorite" and notebook_id:
        store.toggle_favorite(notebook_id)
        st.rerun()
    elif action == "edit" and notebook_id:
        st.session_state["edit_notebook_id"] = notebook_id
        st.query_params["view"] = "create"
        st.rerun()
    elif action == "delete" and notebook_id:
        st.session_state[f"confirm_del_{notebook_id}"] = True

    if view == "notebook":
        nb_id = st.session_state.get("current_notebook_id")
        if not nb_id:
            st.warning("No notebook selected. Returning to list.")
            st.query_params["view"] = "list"
            st.rerun()
            return
        nb = store.get_notebook(nb_id)
        if not nb:
            st.warning("Notebook not found. Returning to list.")
            st.query_params["view"] = "list"
            st.rerun()
            return

        # Header
        header_cols = st.columns([9,1])
        with header_cols[0]:
            st.title(nb.name)
            st.caption(nb.description or "")
            # Info chips: sources count â€¢ tags â€¢ created date
            created_date = (nb.created_at or "").split("T")[0]
            sources_count = len(nb.sources) if nb.sources else 0
            tags_text = ", ".join(nb.tags[:5]) if getattr(nb, 'tags', None) else "No tags"
            st.markdown(
                f"**ðŸ“„ {sources_count} sources** â€¢ **ðŸ·ï¸ {tags_text}** â€¢ **ðŸ“… {created_date}**"
            )
        with header_cols[1]:
            if st.button("ðŸ““ " + t("your_notebooks", _get_lang())):
                st.query_params["view"] = "list"
                st.rerun()

        # Three-tab layout: Notebook | Studio | Source (Settings moved under Studio)
        tab_notebook, tab_studio, tab_sources = st.tabs([t("tab_notebook", _get_lang()), t("tab_studio", _get_lang()), t("tab_sources", _get_lang())])

        with tab_notebook:
            _notebook_overview(nb)
            st.divider()
            _render_chat(nb)

        with tab_studio:
            with st.expander(t("studio_section", _get_lang()), expanded=True):
                # Compact 3x2 grid: Generate | Download (adjacent)
                # Row 1: DOCX
                r1c1, r1c2 = st.columns([10, 1])
                with r1c1:
                    _render_studio_actions(nb, layout="docx_only")
                with r1c2:
                    latest_docx = _get_latest_docx_path(nb.id)
                    if latest_docx.exists():
                        with open(latest_docx, "rb") as f:
                            st.download_button("â¬‡ï¸", f, file_name=latest_docx.name, help=t("download_docx_help", _get_lang()))

                # Row 2: Audio
                r2c1, r2c2 = st.columns([10, 1])
                with r2c1:
                    _render_studio_actions(nb, layout="audio_only")
                with r2c2:
                    latest_audio = _get_latest_audio_path(nb.id)
                    if latest_audio.exists():
                        with open(latest_audio, "rb") as f:
                            st.download_button("â¬‡ï¸", f, file_name=latest_audio.name, help=t("download_audio_help", _get_lang()))

                # Row 3: Mindmap (interactive viewer with background generation)
                r3c1, r3c2 = st.columns([10, 1])
                with r3c1:
                    task_map_key = f"studio_mindmap_task_{nb.id}"
                    if task_map_key not in _TASK_STATUS:
                        _TASK_STATUS[task_map_key] = {"running": False, "file_path": None, "is_html": False, "error": None}
                    map_state = _TASK_STATUS.get(task_map_key, {"running": False})
                    map_label = t("generating_audio", _get_lang()).replace("audio", "mindmap") if map_state.get("running") else t("btn_mindmap", _get_lang())
                    if st.button(map_label, key=f"btn_mindmap_{nb.id}", disabled=map_state.get("running", False)):
                        _TASK_STATUS[task_map_key] = {"running": True, "file_path": None, "is_html": False, "error": None}
                        threading.Thread(target=_background_generate_mindmap, args=(nb.id, task_map_key), daemon=True).start()
                    if map_state.get("error"):
                        st.error(f"âŒ Lá»—i: {map_state['error']}")
                    # Inline viewer if exists
                    html_path = _get_latest_mindmap_html_path(nb.id)
                    png_path = _get_latest_mindmap_path(nb.id)
                    dot_path = _get_notebook_folder(nb.id) / "mindmap_latest.dot"
                    with st.expander("ðŸŽ¥ " + t("open_mindmap", _get_lang()), expanded=True):
                        if html_path.exists():
                            try:
                                with open(html_path, "r", encoding="utf-8") as f:
                                    components.html(f.read(), height=650, width=900, scrolling=True)
                            except Exception:
                                pass
                        elif _is_valid_image(png_path):
                            st.image(str(png_path), caption=t("btn_mindmap", _get_lang()), use_container_width=True)
                        elif dot_path.exists():
                            try:
                                st.graphviz_chart(dot_path.read_text(encoding="utf-8"))
                            except Exception:
                                pass
                with r3c2:
                    # Open mindmap: prefer interactive HTML, fallback to PNG download
                    html_path = _get_latest_mindmap_html_path(nb.id)
                    png_path = _get_latest_mindmap_path(nb.id)

                    if html_path.exists():
                        try:
                            with open(html_path, "r", encoding="utf-8") as f:
                                st.download_button("â¬‡ï¸", f, file_name=html_path.name, help=t("open_mindmap", _get_lang()) + " (HTML)")
                        except Exception:
                            pass
                    elif _is_valid_image(png_path):
                        with open(png_path, "rb") as f:
                            st.download_button("â¬‡ï¸", f, file_name=png_path.name, help=t("download_mindmap_help", _get_lang()) + " (PNG)")

            # Notes section (renamed from Studio inner content)
            with st.expander(t("notes_section", _get_lang()), expanded=True):
                _render_studio_panel(nb)
            st.divider()
            with st.expander(t("settings_section", _get_lang()), expanded=False):
                with st.form("notebook_settings_form"):
                    new_name = st.text_input(t("rename_notebook", _get_lang()), value=nb.name)
                    new_tags_str = st.text_input(t("edit_tags", _get_lang()), value=", ".join(nb.tags) if getattr(nb, 'tags', None) else "")
                    submitted_settings = st.form_submit_button(t("save_settings", _get_lang()), type="primary")
                if submitted_settings:
                    if not new_name.strip():
                        st.error(t("invalid_notebook_name", _get_lang()))
                    else:
                        new_tags_list = [t.strip() for t in new_tags_str.split(',') if t.strip()]
                        store.update_notebook(
                            nb.id,
                            name=new_name.strip(),
                            description=nb.description or "",
                            tags=new_tags_list,
                        )
                        st.success(t("notebook_settings_saved", _get_lang()))
                        st.rerun()

        with tab_sources:
            _render_sources_panel(nb)

        return

    if view == "create":
        _render_create_view()
        return

    lang = _get_lang()
    st.title(t("page_notebooks_title", lang))
    st.markdown(t("page_notebooks_subtitle", lang))

    create_col = st.container()
    with create_col:
        if st.button(t("create_new_notebook", lang), type="primary"):
            st.query_params["view"] = "create"
            st.rerun()

        st.markdown("---")
    st.subheader(t("your_notebooks", lang))

    # Lazy load filters and notebooks
    with st.spinner(t("loading_notebooks", lang)):
        # Get current filters and sorting options
        q, fav, dfrom, dto, sort_by = _render_filters_and_sorting()
        
        # Check if filters or sorting changed and clear cache if needed
        current_filters = f"{q}_{fav}_{dfrom}_{dto}_{sort_by}"
        previous_filters = st.session_state.get('previous_filters', current_filters)
        if previous_filters != current_filters:
            # Clear all notebook caches when filters or sorting change
            for key in list(st.session_state.keys()):
                if key.startswith('notebooks_cache_'):
                    del st.session_state[key]
            st.session_state['previous_filters'] = current_filters
            st.session_state['notebooks_page'] = 0  # Reset to first page
        
        # Cache notebooks in session state to avoid reloading
        cache_key = f"notebooks_cache_{hash(f'{q}_{fav}_{dfrom}_{dto}_{sort_by}')}"
        if cache_key not in st.session_state:
            raw_notebooks = store.list_notebooks(q, fav, dfrom, dto)
            # Apply additional sorting based on user preference
            if sort_by == "date_old":
                raw_notebooks.sort(key=lambda n: n.created_at or "", reverse=False)
            elif sort_by == "updated":
                raw_notebooks.sort(key=lambda n: n.updated_at or n.created_at or "", reverse=True)
            elif sort_by == "name_az":
                raw_notebooks.sort(key=lambda n: n.name.lower())
            elif sort_by == "name_za":
                raw_notebooks.sort(key=lambda n: n.name.lower(), reverse=True)
            # Default "date_new" is already handled by store.list_notebooks
            
            st.session_state[cache_key] = raw_notebooks
        
        notebooks = st.session_state[cache_key]
        
        if not notebooks:
            st.info(t("no_notebooks_yet", lang))
            return
        
        # Lazy render notebooks with pagination
        page_size = 8  # Show 8 notebooks per page
        current_page = st.session_state.get('notebooks_page', 0)
        total_pages = (len(notebooks) + page_size - 1) // page_size
        
        # Pagination controls
        if total_pages > 1:
            col1, col2, col3 = st.columns([1, 2, 1])
            with col1:
                if st.button(t("prev", _get_lang()), disabled=current_page == 0):
                    st.session_state['notebooks_page'] = max(0, current_page - 1)
                    st.rerun()
            with col2:
                st.write(t("page_of", _get_lang()).format(cur=current_page + 1, total=total_pages))
            with col3:
                if st.button(t("next", _get_lang()), disabled=current_page >= total_pages - 1):
                    st.session_state['notebooks_page'] = min(total_pages - 1, current_page + 1)
                    st.rerun()
        
        # Calculate start and end indices for current page
        start_idx = current_page * page_size
        end_idx = min(start_idx + page_size, len(notebooks))
        current_notebooks = notebooks[start_idx:end_idx]
        
        # Render notebooks in grid with lazy loading
        cols = st.columns(4)
        for i, nb in enumerate(current_notebooks):
            col_index = i % 4
            with cols[col_index]:
                # Use container to isolate each notebook card
                with st.container():
                    _render_notebook_card(nb, cols[col_index])
        
        # Show total count
        st.caption(t("showing_range", _get_lang()).format(start=start_idx + 1, end=end_idx, total=len(notebooks)))
        
        # Clear cache button for debugging (disable during searching)
        if st.button(t("refresh_notebooks", lang), help=t("refresh_notebooks_help", lang), disabled=st.session_state.get('is_searching', False)):
            # Clear all notebook caches
            for key in list(st.session_state.keys()):
                if key.startswith('notebooks_cache_'):
                    del st.session_state[key]
            st.session_state['notebooks_page'] = 0
            st.rerun()


if __name__ == "__main__":
    main()


